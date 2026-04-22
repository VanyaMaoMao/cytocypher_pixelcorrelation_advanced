from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Dict, List, Optional, Sequence

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from scipy.signal import peak_prominences, peak_widths

from .plotting import plot_afc_review_item
from .results import AFCEvent, AFCSegmentReviewItem, _format_events_df_for_report, _format_rescue_df_for_report

def _add_dataframe_to_docx(doc: Document, df: pd.DataFrame, empty_text: str = "No detected events.") -> None:
    if df.empty:
        doc.add_paragraph(str(empty_text))
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.autofit = True
    for j, c in enumerate(df.columns):
        table.cell(0, j).text = str(c)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            cells[j].text = "" if pd.isna(v) else str(v)


def _as_float(v, default: float = np.nan) -> float:
    try:
        vf = float(v)
    except Exception:
        return float(default)
    return float(vf) if np.isfinite(vf) else float(default)


def _sample_id_as_value(v):
    if v is None:
        return np.nan
    if isinstance(v, str):
        s = v.strip()
        if not s:
            return np.nan
        try:
            vf = float(s)
            if np.isfinite(vf) and abs(vf - round(vf)) < 1e-9:
                return int(round(vf))
        except Exception:
            pass
        return s
    try:
        vf = float(v)
        if np.isfinite(vf):
            if abs(vf - round(vf)) < 1e-9:
                return int(round(vf))
            return vf
    except Exception:
        pass
    return v


def _sample_id_as_text(v) -> str:
    val = _sample_id_as_value(v)
    if pd.isna(val):
        return ""
    return str(val)


def _prettify_table_headers(df: pd.DataFrame, rename_map: Optional[Dict[str, str]] = None) -> pd.DataFrame:
    out = df.copy()
    rename_map = rename_map or {}
    pretty_cols: Dict[str, str] = {}
    for c in out.columns:
        if c in rename_map:
            pretty_cols[c] = str(rename_map[c])
            continue
        label = str(c).replace("_", " ").strip().title()
        if label.endswith(" S"):
            label = label[:-2] + " (s)"
        pretty_cols[c] = label
    return out.rename(columns=pretty_cols)


def _format_float_columns(df: pd.DataFrame, cols: Sequence[str], ndigits: int) -> pd.DataFrame:
    out = df.copy()
    for c in cols:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce").map(
                lambda v: "" if pd.isna(v) else f"{float(v):.{int(ndigits)}f}"
            )
    return out


def _segment_afc_review_row(review_log_df: Optional[pd.DataFrame], seg_no: int) -> Optional[pd.Series]:
    if review_log_df is None or review_log_df.empty or "segment_index" not in review_log_df.columns:
        return None
    idx = pd.to_numeric(review_log_df["segment_index"], errors="coerce")
    seg_rows = review_log_df.loc[idx == int(seg_no)]
    if seg_rows.empty:
        return None
    return seg_rows.iloc[0]


def _main_peaks_from_events(events_df: pd.DataFrame) -> tuple[list[float], list[float]]:
    if events_df is None or events_df.empty or "Type" not in events_df.columns:
        return [], []
    m = events_df.loc[events_df["Type"].astype(str) == "Main Beat"].copy()
    if m.empty:
        return [], []
    m["Time_s"] = pd.to_numeric(m["Time_s"], errors="coerce")
    m["Amp"] = pd.to_numeric(m.get("Amp", np.nan), errors="coerce")
    m = m.dropna(subset=["Time_s"]).sort_values("Time_s").reset_index(drop=True)
    return [float(x) for x in m["Time_s"].tolist()], [float(x) for x in m["Amp"].tolist()]


def _rescue_peaks_from_meta(meta: Dict, time: np.ndarray, sig: np.ndarray) -> tuple[list[float], list[float]]:
    rescue_idx = np.asarray(meta.get("_rescue_peaks_plot", []), dtype=int)
    rescue_idx = rescue_idx[(rescue_idx >= 0) & (rescue_idx < len(time))]
    if rescue_idx.size == 0:
        return [], []
    return [float(time[i]) for i in rescue_idx.tolist()], [float(sig[i]) for i in rescue_idx.tolist()]


def _rescue_events_df_from_meta(meta: Dict, time: np.ndarray, sig: np.ndarray) -> pd.DataFrame:
    rescue_idx = np.asarray(meta.get("_rescue_peaks_plot", []), dtype=int)
    rescue_idx = rescue_idx[(rescue_idx >= 0) & (rescue_idx < len(time))]
    if rescue_idx.size == 0:
        return pd.DataFrame(columns=["Time_s", "Amp", "Prom", "Width_s"])
    s = np.asarray(sig, dtype=float)
    t = np.asarray(time, dtype=float)
    if t.size > 1:
        fs = float(1.0 / max(np.median(np.diff(t)), 1e-9))
    else:
        fs = np.nan
    prom = peak_prominences(s, rescue_idx)[0] if rescue_idx.size else np.array([], dtype=float)
    width_samples = peak_widths(s, rescue_idx, rel_height=0.5)[0] if rescue_idx.size else np.array([], dtype=float)
    widths = (width_samples / fs) if np.isfinite(fs) and fs > 0 else np.full_like(width_samples, np.nan, dtype=float)
    return pd.DataFrame(
        {
            "Time_s": [float(t[i]) for i in rescue_idx.tolist()],
            "Amp": [float(s[i]) for i in rescue_idx.tolist()],
            "Prom": [float(x) for x in np.asarray(prom, dtype=float).tolist()],
            "Width_s": [float(x) for x in np.asarray(widths, dtype=float).tolist()],
        }
    ).sort_values("Time_s").reset_index(drop=True)


def _clean_afc_review_png_bytes(
    *,
    sheet_name: str,
    seg_no: int,
    meta: Dict,
    events_df: pd.DataFrame,
    seg_afc_df: pd.DataFrame,
    review_row: Optional[pd.Series],
) -> Optional[bytes]:
    time = np.asarray(meta.get("_time_plot", []), dtype=float)
    sig = np.asarray(meta.get("_sig_plot", []), dtype=float)
    if time.size == 0 or sig.size == 0 or time.size != sig.size:
        return None

    x_start = float(time[0])
    x_end = float(time[-1])
    low_l = np.nan
    low_r = np.nan
    up_l = np.nan
    up_r = np.nan
    status = "not_reviewed"
    if review_row is not None:
        x_start = _as_float(review_row.get("x_start_s", x_start), x_start)
        x_end = _as_float(review_row.get("x_end_s", x_end), x_end)
        low_l = _as_float(review_row.get("afc_lower_left_value", np.nan), np.nan)
        low_r = _as_float(review_row.get("afc_lower_right_value", np.nan), np.nan)
        up_l = _as_float(review_row.get("afc_upper_left_value", np.nan), np.nan)
        up_r = _as_float(review_row.get("afc_upper_right_value", np.nan), np.nan)
        status = str(review_row.get("status", status))

    main_t, main_a = _main_peaks_from_events(events_df)
    rescue_t, rescue_a = _rescue_peaks_from_meta(meta, time, sig)
    man_t = []
    man_a = []
    if seg_afc_df is not None and not seg_afc_df.empty:
        man_t = [float(x) for x in pd.to_numeric(seg_afc_df["time_s"], errors="coerce").dropna().tolist()]
        man_a = [float(x) for x in pd.to_numeric(seg_afc_df["amplitude"], errors="coerce").dropna().tolist()]

    item = AFCSegmentReviewItem(
        segment_name=str(sheet_name),
        segment_index=int(seg_no),
        afc_lower_left_value=float(low_l),
        afc_lower_right_value=float(low_r),
        afc_upper_left_value=float(up_l),
        afc_upper_right_value=float(up_r),
        x_start_s=float(x_start),
        x_end_s=float(x_end),
        status=str(status),
        main_peak_times_s=main_t,
        main_peak_amps=main_a,
        rescue_peak_times_s=rescue_t,
        rescue_peak_amps=rescue_a,
        helper_candidate_times_s=[],
        helper_candidate_amps=[],
        manual_afc_times_s=man_t,
        manual_afc_amps=man_a,
    )
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    plot_afc_review_item(
        ax=ax,
        time_array=time,
        signal_array=sig,
        review_item=item,
        helper_candidate_times_s=[],
        manual_afc_times_s=man_t,
        title=f"{sheet_name} | AFC review (final curated)",
        show_helper_candidates=False,
    )
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()

def build_raw_cytocypher_docx_report(
    *,
    raw_xlsx_path: str,
    stim_hz: float,
    recording_s: float,
    segment_results: List[Dict],
    output_docx: str,
    afc_events: Optional[Sequence[AFCEvent]] = None,
    review_log_df: Optional[pd.DataFrame] = None,
    afc_review_plots_dir: Optional[str] = None,
) -> None:
    doc = Document()
    doc.add_heading("CytoCypher arrhythmia analysis report", level=0)
    doc.add_paragraph(f"Source file: {os.path.basename(raw_xlsx_path)}")
    if stim_hz > 0:
        doc.add_paragraph(f"Stim frequency (Hz): {stim_hz:.4g}")
    else:
        doc.add_paragraph("Stim frequency (Hz): 0.0 (spontaneous)")
    doc.add_paragraph(f"Recording duration (s): {recording_s:.4g}")
    doc.add_paragraph("Processed sheet types: PixelCorrelation Segment")

    afc_df_all = pd.DataFrame([x.to_dict() for x in afc_events]) if afc_events else pd.DataFrame()
    if not afc_df_all.empty:
        doc.add_heading("AFC Review Summary", level=1)
        doc.add_paragraph(f"Total AFC events: {len(afc_df_all)}")
        by_src = afc_df_all["source"].astype(str).value_counts(dropna=False).to_dict()
        if by_src:
            src_lines = [f"{k}: {int(v)}" for k, v in by_src.items()]
            doc.add_paragraph("AFC source counts:\n" + "\n".join(src_lines))
        if review_log_df is not None and not review_log_df.empty:
            done_n = int((review_log_df["status"].astype(str).str.lower() != "skipped").sum())
            skip_n = int((review_log_df["status"].astype(str).str.lower() == "skipped").sum())
            doc.add_paragraph(f"Reviewed segments: {done_n} | skipped: {skip_n}")

    for item in segment_results:
        sheet_name = str(item["sheet_name"])
        meta = item["meta"]
        events = item["events"]
        sample_id = _sample_id_as_text(item.get("sample_id", meta.get("sample_id", np.nan)))
        n_total = int(item["n_main"])
        n_primary = int(item.get("n_main_primary", n_total))
        n_rescue = int(item.get("n_rescue", 0))
        bpm_user = float((n_total / recording_s) * 60.0) if recording_s > 0 else 0.0

        heading = f"{sheet_name} (Sample ID {sample_id})" if sample_id else sheet_name
        doc.add_heading(heading, level=1)
        doc.add_paragraph(f"QC: {'PASS' if bool(meta.get('qc_pass', False)) else 'REJECT'} ({meta.get('qc_reason', 'unknown')})")

        png_bytes = item.get("plot_png", None)
        if isinstance(png_bytes, (bytes, bytearray)) and len(png_bytes) > 0:
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.8))

        metric_lines = [
            f"BPM (using user duration): {bpm_user:.2f}",
            f"Primary main beats: {n_primary}",
            f"Rescue peaks: {n_rescue}",
            f"Total detected events: {n_total}",
        ]
        if stim_hz > 0 and recording_s > 0:
            expected = float(stim_hz * recording_s)
            ratio = float(n_total / expected) if expected > 0 else np.nan
            metric_lines.append(f"Expected beats (Hz x seconds): {expected:.3f}")
            metric_lines.append(
                f"Detected Total/Expected ratio: {ratio:.3f}" if np.isfinite(ratio) else "Detected Total/Expected ratio: N/A"
            )
        doc.add_paragraph("\n".join(metric_lines) + "\n")

        doc.add_paragraph("Detected events table")
        _add_dataframe_to_docx(doc, _format_events_df_for_report(events))
        doc.add_paragraph("Rescue peaks (between transients)")
        rescue_df = _rescue_events_df_from_meta(
            meta,
            np.asarray(meta.get("_time_plot", []), dtype=float),
            np.asarray(meta.get("_sig_plot", []), dtype=float),
        )
        _add_dataframe_to_docx(
            doc,
            _format_rescue_df_for_report(rescue_df),
            empty_text="No rescue peaks.",
        )

        if (review_log_df is not None and not review_log_df.empty) or not afc_df_all.empty:
            seg_no = int(item.get("segment_no", -1))
            seg_afc = afc_df_all.loc[afc_df_all["segment_index"].astype(int) == seg_no].copy() if not afc_df_all.empty else pd.DataFrame()
            seg_afc = seg_afc.sort_values("time_s").reset_index(drop=True) if not seg_afc.empty else seg_afc
            review_row = _segment_afc_review_row(review_log_df, seg_no)

            # Mini table 1: segment AFC review settings/status.
            if review_row is None:
                review_table_df = pd.DataFrame(
                    [
                        {
                            "afc_lower_left_value": np.nan,
                            "afc_lower_right_value": np.nan,
                            "afc_upper_left_value": np.nan,
                            "afc_upper_right_value": np.nan,
                            "x_start_s": np.nan,
                            "x_end_s": np.nan,
                            "status": "not_reviewed",
                        }
                    ]
                )
            else:
                review_table_df = pd.DataFrame(
                    [
                        {
                            "afc_lower_left_value": _as_float(review_row.get("afc_lower_left_value", np.nan), np.nan),
                            "afc_lower_right_value": _as_float(review_row.get("afc_lower_right_value", np.nan), np.nan),
                            "afc_upper_left_value": _as_float(review_row.get("afc_upper_left_value", np.nan), np.nan),
                            "afc_upper_right_value": _as_float(review_row.get("afc_upper_right_value", np.nan), np.nan),
                            "x_start_s": _as_float(review_row.get("x_start_s", np.nan), np.nan),
                            "x_end_s": _as_float(review_row.get("x_end_s", np.nan), np.nan),
                            "status": str(review_row.get("status", "pending")),
                        }
                    ]
                )
            review_table_df = _format_float_columns(
                review_table_df,
                cols=[
                    "afc_lower_left_value",
                    "afc_lower_right_value",
                    "afc_upper_left_value",
                    "afc_upper_right_value",
                    "x_start_s",
                    "x_end_s",
                ],
                ndigits=6,
            )
            review_table_df = _prettify_table_headers(
                review_table_df,
                rename_map={
                    "afc_lower_left_value": "AFC Lower Left",
                    "afc_lower_right_value": "AFC Lower Right",
                    "afc_upper_left_value": "AFC Upper Left",
                    "afc_upper_right_value": "AFC Upper Right",
                    "x_start_s": "X Start (s)",
                    "x_end_s": "X End (s)",
                    "status": "Status",
                },
            )
            doc.add_paragraph("AFC review settings")
            _add_dataframe_to_docx(doc, review_table_df)

            # Clean AFC report plot: no helper crosses.
            clean_plot_bytes = _clean_afc_review_png_bytes(
                sheet_name=sheet_name,
                seg_no=seg_no,
                meta=meta,
                events_df=events,
                seg_afc_df=seg_afc,
                review_row=review_row,
            )
            if isinstance(clean_plot_bytes, (bytes, bytearray)) and len(clean_plot_bytes) > 0:
                doc.add_picture(io.BytesIO(clean_plot_bytes), width=Inches(6.8))
            elif afc_review_plots_dir:
                # Fallback: keep legacy behavior if clean plot cannot be generated.
                plot_dir = Path(afc_review_plots_dir)
                if plot_dir.exists():
                    pattern = f"segment_{seg_no:03d}_review.png"
                    seg_plots = sorted(plot_dir.glob(pattern))[:1]
                    for p in seg_plots:
                        try:
                            doc.add_picture(str(p), width=Inches(6.8))
                        except Exception:
                            pass

            # Mini table 2: final AFC events (selected points only).
            doc.add_paragraph("AFC selected events")
            if seg_afc.empty:
                _add_dataframe_to_docx(
                    doc,
                    pd.DataFrame(columns=["Time_s", "Amp", "Prom", "Width_s"]),
                    empty_text="No final AFC events.",
                )
            else:
                seg_events_df = pd.DataFrame(
                    {
                        "Time_s": pd.to_numeric(seg_afc.get("time_s", np.nan), errors="coerce"),
                        "Amp": pd.to_numeric(seg_afc.get("amplitude", np.nan), errors="coerce"),
                        "Prom": pd.to_numeric(seg_afc.get("prominence", np.nan), errors="coerce"),
                        "Width_s": pd.to_numeric(seg_afc.get("width_s", np.nan), errors="coerce"),
                    }
                )
                seg_events_df = _format_float_columns(seg_events_df, cols=["Time_s"], ndigits=3)
                seg_events_df = _format_float_columns(seg_events_df, cols=["Amp", "Prom", "Width_s"], ndigits=5)
                _add_dataframe_to_docx(doc, seg_events_df.reset_index(drop=True))

    doc.save(output_docx)


def _collect_main_events_table(segment_results: Sequence[Dict]) -> pd.DataFrame:
    rows: List[pd.DataFrame] = []
    for i, seg in enumerate(segment_results):
        seg_name = str(seg.get("sheet_name", f"Segment {i + 1}"))
        sample_id = _sample_id_as_value(seg.get("sample_id", np.nan))
        seg_idx = int(seg.get("segment_no", i + 1))

        df_main = seg.get("events")
        if isinstance(df_main, pd.DataFrame) and (not df_main.empty):
            seg_main = df_main.copy()
            for c in ["Time_s", "Type", "Amp", "Prom", "Width_s", "Transient"]:
                if c not in seg_main.columns:
                    seg_main[c] = np.nan
            seg_main = seg_main[["Time_s", "Type", "Amp", "Prom", "Width_s", "Transient"]]
            seg_main.insert(0, "Segment", seg_name)
            seg_main.insert(1, "Sample ID", sample_id)
            seg_main.insert(2, "SegmentIndex", seg_idx)
            rows.append(seg_main)

        meta = dict(seg.get("meta", {}) or {})
        rescue_df = _rescue_events_df_from_meta(
            meta,
            np.asarray(meta.get("_time_plot", []), dtype=float),
            np.asarray(meta.get("_sig_plot", []), dtype=float),
        )
        if not rescue_df.empty:
            rescue_rows = rescue_df.copy()
            rescue_rows["Type"] = "Rescue"
            rescue_rows["Transient"] = np.nan
            rescue_rows = rescue_rows[["Time_s", "Type", "Amp", "Prom", "Width_s", "Transient"]]
            rescue_rows.insert(0, "Segment", seg_name)
            rescue_rows.insert(1, "Sample ID", sample_id)
            rescue_rows.insert(2, "SegmentIndex", seg_idx)
            rows.append(rescue_rows)

    if not rows:
        return pd.DataFrame(columns=["Segment", "Sample ID", "SegmentIndex", "Time_s", "Type", "Amp", "Prom", "Width_s", "Transient"])
    out = pd.concat(rows, ignore_index=True, sort=False)
    out["Time_s"] = pd.to_numeric(out["Time_s"], errors="coerce")
    return out.sort_values(["SegmentIndex", "Time_s"]).reset_index(drop=True)


def _collect_afc_events_table(afc_events: Optional[Sequence[AFCEvent]], segment_results: Sequence[Dict]) -> pd.DataFrame:
    by_seg_idx: Dict[int, object] = {}
    by_seg_name: Dict[str, object] = {}
    for i, seg in enumerate(segment_results):
        seg_idx = int(seg.get("segment_no", i + 1))
        seg_name = str(seg.get("sheet_name", f"Segment {seg_idx}"))
        sid = _sample_id_as_value(seg.get("sample_id", seg.get("meta", {}).get("sample_id", np.nan)))
        by_seg_idx[seg_idx] = sid
        by_seg_name[seg_name] = sid

    cols = [
        "segment_name",
        "sample_id",
        "segment_index",
        "main_peak_index",
        "time_s",
        "amplitude",
        "prominence",
        "width_s",
        "review_id",
    ]
    if not afc_events:
        return pd.DataFrame(columns=cols)
    df = pd.DataFrame([x.to_dict() for x in afc_events])
    if df.empty:
        return pd.DataFrame(columns=cols)
    sid_vals: List[object] = []
    for _, r in df.iterrows():
        seg_idx = pd.to_numeric(r.get("segment_index", np.nan), errors="coerce")
        seg_name = str(r.get("segment_name", ""))
        sid = np.nan
        if pd.notna(seg_idx):
            sid = by_seg_idx.get(int(seg_idx), np.nan)
        if pd.isna(sid):
            sid = by_seg_name.get(seg_name, np.nan)
        sid_vals.append(_sample_id_as_value(sid))
    df["sample_id"] = sid_vals
    for c in cols:
        if c not in df.columns:
            df[c] = np.nan
    return df[cols].sort_values(["segment_index", "main_peak_index", "time_s"]).reset_index(drop=True)


def _build_clean_summary_sheet(
    summary_df: pd.DataFrame,
    main_events_df: pd.DataFrame,
    segment_results: Sequence[Dict],
) -> pd.DataFrame:
    sample_id_map: Dict[str, object] = {}
    for i, seg in enumerate(segment_results):
        seg_name = str(seg.get("sheet_name", f"Segment {i + 1}"))
        sample_id_map[seg_name] = _sample_id_as_value(seg.get("sample_id", seg.get("meta", {}).get("sample_id", np.nan)))

    cols = [
        "Segment",
        "Sample ID",
        "QC",
        "QC reason",
        "BPM (using user duration)",
        "Primary main beats",
        "Rescue peaks",
        "Total detected events",
        "Expected beats (Hz x seconds)",
        "Detected Total/Expected ratio",
        "detected_total_expected_ratio",
        "Average Prom",
        "Average AMP",
    ]
    if summary_df is None or summary_df.empty:
        return pd.DataFrame(columns=cols)

    n_rows = len(summary_df)

    def _pick_col(*names: str, default: object = np.nan) -> pd.Series:
        for nm in names:
            if nm in summary_df.columns:
                return summary_df[nm]
        return pd.Series([default] * n_rows)

    out = pd.DataFrame(
        {
            "Segment": _pick_col("Segment", default=""),
            "Sample ID": np.nan,
            "QC": _pick_col("QC", default=""),
            "QC reason": _pick_col("QC reason", default=""),
            "BPM (using user duration)": _pick_col("BPM (using user duration)", default=np.nan),
            "Primary main beats": _pick_col("Primary main beats", "primary_main_beats", default=np.nan),
            "Rescue peaks": _pick_col("Rescue peaks", "rescue_peaks", default=np.nan),
            "Total detected events": _pick_col("Total detected events", "total_detected_events", "Main beats", default=np.nan),
            "Expected beats (Hz x seconds)": _pick_col("Expected beats (Hz x seconds)", "expected_beats", default=np.nan),
            "Detected Total/Expected ratio": _pick_col(
                "Detected Total/Expected ratio",
                "detected_total_expected_ratio",
                "Detected Main/Expected ratio",
                default=np.nan,
            ),
            "detected_total_expected_ratio": _pick_col(
                "detected_total_expected_ratio",
                "Detected Total/Expected ratio",
                "Detected Main/Expected ratio",
                default=np.nan,
            ),
        }
    )
    out["Sample ID"] = out["Segment"].map(lambda s: _sample_id_as_value(sample_id_map.get(str(s), np.nan)))

    avg_df = pd.DataFrame(columns=["Segment", "Average Prom", "Average AMP"])
    if main_events_df is not None and not main_events_df.empty:
        mm = main_events_df.copy()
        mm["Prom"] = pd.to_numeric(mm.get("Prom", np.nan), errors="coerce")
        mm["Amp"] = pd.to_numeric(mm.get("Amp", np.nan), errors="coerce")
        avg_df = (
            mm.groupby("Segment", as_index=False)
            .agg(**{"Average Prom": ("Prom", "mean"), "Average AMP": ("Amp", "mean")})
        )

    out = out.merge(avg_df, on="Segment", how="left")
    for c in ["Average Prom", "Average AMP"]:
        if c not in out.columns:
            out[c] = np.nan

    reject_mask = out["QC"].astype(str).str.upper().eq("REJECT")
    for c in [
        "BPM (using user duration)",
        "Primary main beats",
        "Rescue peaks",
        "Total detected events",
        "Detected Total/Expected ratio",
        "detected_total_expected_ratio",
        "Average Prom",
        "Average AMP",
    ]:
        out.loc[reject_mask, c] = np.nan

    return out[cols]


def build_arrhythmia_summary_workbook(
    *,
    output_xlsx: str,
    summary_df: pd.DataFrame,
    segment_results: Sequence[Dict],
    afc_events: Optional[Sequence[AFCEvent]] = None,
    review_log_df: Optional[pd.DataFrame] = None,
) -> None:
    main_events_df = _collect_main_events_table(segment_results)
    afc_events_df = _collect_afc_events_table(afc_events, segment_results)
    summary_clean_df = _build_clean_summary_sheet(summary_df, main_events_df, segment_results)
    if review_log_df is None or (review_log_df.empty and len(review_log_df.columns) == 0):
        review_log_df = pd.DataFrame(
            columns=[
                "segment_name",
                "segment_index",
                "afc_lower_left_value",
                "afc_lower_right_value",
                "afc_upper_left_value",
                "afc_upper_right_value",
                "x_start_s",
                "x_end_s",
                "status",
            ]
        )
    review_log_public = review_log_df.drop(columns=["manual_afc_times_s", "manual_afc_amps"], errors="ignore").copy()

    with pd.ExcelWriter(output_xlsx, engine="openpyxl") as writer:
        summary_clean_df.to_excel(writer, index=False, sheet_name="summary")
        main_events_df.to_excel(writer, index=False, sheet_name="main_events")
        afc_events_df.to_excel(writer, index=False, sheet_name="afc_events")
        review_log_public.to_excel(writer, index=False, sheet_name="review_log")

